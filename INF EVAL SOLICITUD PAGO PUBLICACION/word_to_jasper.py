import sys
import os
import re as regex
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree as ET

EMU_PER_POINT = 914400
EMU_PER_PIXEL = 9142

NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'wps': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingShape',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    'wps2': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
}


class JasperConverter:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.docx_dir = os.path.dirname(docx_path)
        self.doc = Document(docx_path)
        self.report_name = os.path.splitext(os.path.basename(docx_path))[0].replace(" ", "_")

        self.images_dir = os.path.join(os.path.dirname(docx_path), f".images_{self.report_name}")
        os.makedirs(self.images_dir, exist_ok=True)

        self.header_image_map = {}
        self.header_images_copied = self._extract_header_images()

        self.current_band_type = "detail"
        self.ns_uri = "http://jasperreports.sourceforge.net/jasperreports"
        self.jrxml, self.bands = self._init_xml()

    def _init_xml(self):
        nsmap = {None: self.ns_uri, "xsi": "http://www.w3.org/2001/XMLSchema-instance"}
        root = ET.Element("jasperReport", nsmap=nsmap)
        root.set("name", self.report_name)
        root.set("whenNoDataType", "AllSectionsNoDetail")
        root.set("{http://www.w3.org/2001/XMLSchema-instance}schemaLocation",
                  f"{self.ns_uri} http://jasperreports.sourceforge.net/xsd/jasperreport.xsd")

        bands_map = {}
        for section in ["background", "title", "pageHeader", "columnHeader", "detail", "columnFooter", "pageFooter", "summary"]:
            bands_map[section] = ET.SubElement(root, section)

        return root, bands_map

    def emu_to_pt(self, emu_value):
        if emu_value is None:
            return 0
        if hasattr(emu_value, 'pt'):
            return emu_value.pt
        try:
            return float(emu_value) / EMU_PER_POINT
        except:
            return 0

    def emu_docx_to_pt(self, emu_value):
        if emu_value is None:
            return 0
        try:
            return float(emu_value) / 12700
        except:
            return 0

    def get_alignment(self, para):
        align = para.alignment
        if align == WD_ALIGN_PARAGRAPH.CENTER:
            return "Center"
        if align == WD_ALIGN_PARAGRAPH.RIGHT:
            return "Right"
        if align == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "Justified"
        return "Left"

    def get_styled_text(self, para):
        if not para.runs:
            clean_text = para.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return f'<style size="10.0" fontName="Arial">{clean_text}</style>'

        styled_content = ""
        for run in para.runs:
            font_size_pt = self.emu_to_pt(run.font.size) if run.font.size else 10.0
            f_name = f' fontName="{run.font.name}"' if run.font.name else ' fontName="Arial"'

            styles = []
            if run.bold:
                styles.append('isBold="true"')
            if run.italic:
                styles.append('isItalic="true"')

            style_str = " ".join(styles)
            styled = f'<style size="{font_size_pt}"{f_name}{" " + style_str if styles else ""}>'

            clean_text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            styled_content += f"{styled}{clean_text}</style>"

        return styled_content

    def _extract_header_images(self):
        header_xml_path = os.path.join(self.docx_dir, "word", "header1.xml")
        header_rels_path = os.path.join(self.docx_dir, "word", "_rels", "header1.xml.rels")

        if not os.path.exists(header_xml_path) or not os.path.exists(header_rels_path):
            return {}

        try:
            rels_tree = ET.parse(header_rels_path)
            ns_rels = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}

            for rel in rels_tree.findall('.//r:Relationship', ns_rels):
                r_id = rel.get('Id')
                target = rel.get('Target')
                if target and r_id:
                    src_path = os.path.join(self.docx_dir, "word", target)
                    if os.path.exists(src_path):
                        ext = os.path.splitext(target)[1]
                        dest_name = f"header_{r_id}{ext}"
                        dest_path = os.path.join(self.images_dir, dest_name)
                        shutil.copy2(src_path, dest_path)
                        self.header_image_map[r_id] = dest_path
            return self.header_image_map
        except Exception as e:
            print(f"Error extracting header images: {e}")
            return {}

    def _add_image_to_band(self, band, img_path, x, y, width, height):
        image_elem = ET.SubElement(band, "image")

        re = ET.SubElement(image_elem, "reportElement")
        re.set("x", str(x))
        re.set("y", str(y))
        re.set("width", str(width))
        re.set("height", str(height))

        image_expr = ET.SubElement(image_elem, "imageExpression")
        image_expr.text = f'"{img_path}"'

    def _get_cell_color(self, tc):
        tc_pr = tc.find('.//w:tcPr', NS)
        if tc_pr is not None:
            shd = tc_pr.find('.//w:shd', NS)
            if shd is not None:
                fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill:
                    return f"#{fill}"
        return None

    def _get_text_from_cell(self, cell_elem):
        text_parts = []
        for para in cell_elem.findall('.//w:p', NS):
            para_text = ''.join(
                ''.join(r.text or '' for r in run.findall('.//w:t', NS))
                for run in para.findall('.//w:r', NS)
            )
            if para_text:
                text_parts.append(para_text)
        return ' '.join(text_parts)

    def _get_table_grid_from_docx(self, table):
        if hasattr(table, 'columns') and table.columns:
            cols = []
            for col in table.columns:
                width_emu = col.width if hasattr(col, 'width') and col.width else 914400
                cols.append(int(width_emu))
            return cols
        return []

    def _emu_cols_to_pt(self, col_widths_emu):
        return [self.emu_docx_to_pt(w) for w in col_widths_emu]

    def _process_header_simple(self, header_band):
        header_xml_path = os.path.join(self.docx_dir, "word", "header1.xml")
        if not os.path.exists(header_xml_path):
            for para in self.doc.sections[0].header.paragraphs[:3]:
                text = para.text.strip()
                if text:
                    static = ET.SubElement(header_band, "staticText")
                    re = ET.SubElement(static, "reportElement")
                    re.set("x", "0")
                    re.set("y", str(len(list(header_band)) * 15))
                    re.set("width", "500")
                    re.set("height", "15")
                    te = ET.SubElement(static, "textElement")
                    te.set("textAlignment", "Center")
                    st = ET.SubElement(static, "text")
                    st.text = f'<style size="10" fontName="Arial" isBold="true">{text}</style>'
            return

        try:
            tree = ET.parse(header_xml_path)
            root = tree.getroot()

            anchor = root.find('.//wp:anchor', NS)
            if anchor is not None:
                inline = anchor.find('.//wp:inline', NS)
                if inline is not None:
                    pic = inline.find('.//pic:pic', NS)
                    if pic is not None:
                        blip = pic.find('.//a:blip', NS)
                        if blip is not None:
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed and embed in self.header_image_map:
                                try:
                                    from PIL import Image
                                    with Image.open(self.header_image_map[embed]) as img:
                                        img_width, img_height = img.size
                                        self._add_image_to_band(header_band, self.header_image_map[embed], 0, 0, img_width, img_height)
                                except:
                                    self._add_image_to_band(header_band, self.header_image_map[embed], 0, 0, 100, 30)

            for para in self.doc.sections[0].header.paragraphs:
                text = para.text.strip()
                if text and len(text) > 2:
                    static = ET.SubElement(header_band, "staticText")
                    re = ET.SubElement(static, "reportElement")
                    re.set("x", "100")
                    re.set("y", str(len(list(header_band)) * 15))
                    re.set("width", "400")
                    re.set("height", "15")
                    te = ET.SubElement(static, "textElement")
                    te.set("textAlignment", "Center")
                    st = ET.SubElement(static, "text")
                    st.text = f'<style size="9" fontName="Arial">{text}</style>'

        except Exception as e:
            print(f"Error processing header: {e}")

    def _convert_table_simple(self, table, parent_band, start_y=0):
        col_widths_emu = self._get_table_grid_from_docx(table)
        if not col_widths_emu:
            return

        col_widths = self._emu_cols_to_pt(col_widths_emu)
        total_cols = len(col_widths)

        row_idx = 0
        for row in table.rows:
            cells = row.cells
            x_pos = 0

            for cell_idx, cell in enumerate(cells):
                if cell_idx >= total_cols:
                    break

                cell_width = int(col_widths[cell_idx])
                cell_text = cell.text.strip()
                cell_color = None

                try:
                    cell_xml = cell._element
                    cell_color = self._get_cell_color(cell_xml)
                except:
                    pass

                tf = ET.SubElement(parent_band, "textField")

                re = ET.SubElement(tf, "reportElement")
                re.set("x", str(x_pos))
                re.set("y", str(start_y + row_idx * 20))
                re.set("width", str(cell_width))
                re.set("height", "20")

                box = ET.SubElement(tf, "box")
                pen = ET.SubElement(box, "pen")
                pen.set("lineWidth", "0.5")
                top = ET.SubElement(box, "topPen")
                top.set("lineWidth", "0.5")
                left = ET.SubElement(box, "leftPen")
                left.set("lineWidth", "0.5")
                bottom = ET.SubElement(box, "bottomPen")
                bottom.set("lineWidth", "0.5")
                right = ET.SubElement(box, "rightPen")
                right.set("lineWidth", "0.5")

                if cell_color and cell_color != "#auto":
                    re.set("backcolor", cell_color)
                    re.set("mode", "Opaque")

                te = ET.SubElement(tf, "textElement")
                te.set("verticalAlignment", "Middle")

                if cell_idx == 0:
                    te.set("textAlignment", "Center")

                if cell_text:
                    exp = ET.SubElement(tf, "textFieldExpression")
                    exp.text = f'<style size="8" fontName="Arial">{"<![CDATA[" + cell_text + "]]>"}</style>'

                x_pos += cell_width

            row_idx += 1

        return row_idx * 20

    def add_paragraph(self, para, band, y_offset):
        text = para.text.strip()
        if not text:
            return y_offset

        text_field = ET.SubElement(band, "textField")

        report_elt = ET.SubElement(text_field, "reportElement")
        report_elt.set("x", "0")
        report_elt.set("y", str(y_offset))
        report_elt.set("width", "500")
        report_elt.set("height", "20")

        text_elt = ET.SubElement(text_field, "textElement")
        text_elt.set("textAlignment", self.get_alignment(para))
        text_elt.set("markup", "styled")

        exp = ET.SubElement(text_field, "textFieldExpression")
        exp.text = f'"{self.get_styled_text(para)}"'

        return y_offset + 20

    def process(self):
        section = self.doc.sections[0]
        page_width = int(self.emu_to_pt(section.page_width))
        col_width = int(page_width - self.emu_to_pt(section.left_margin) - self.emu_to_pt(section.right_margin))

        self.jrxml.set("pageWidth", str(page_width))
        self.jrxml.set("pageHeight", str(int(self.emu_to_pt(section.page_height))))
        self.jrxml.set("columnWidth", str(col_width))
        self.jrxml.set("leftMargin", str(int(self.emu_to_pt(section.left_margin))))
        self.jrxml.set("rightMargin", str(int(self.emu_to_pt(section.right_margin))))
        self.jrxml.set("topMargin", str(int(self.emu_to_pt(section.top_margin))))
        self.jrxml.set("bottomMargin", str(int(self.emu_to_pt(section.bottom_margin))))

        header_section = self.bands["pageHeader"]
        header_band = ET.SubElement(header_section, "band")
        header_band.set("height", "60")
        self._process_header_simple(header_band)

        footer_section = self.bands["pageFooter"]
        footer_band = ET.SubElement(footer_section, "band")
        footer_band.set("height", "20")
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                static = ET.SubElement(footer_band, "staticText")
                re = ET.SubElement(static, "reportElement")
                re.set("x", "0")
                re.set("y", "0")
                re.set("width", str(col_width))
                re.set("height", "20")
                te = ET.SubElement(static, "textElement")
                te.set("textAlignment", "Center")
                st = ET.SubElement(static, "text")
                st.text = f'<style size="8" fontName="Arial">{text}</style>'

        current_section = self.bands["detail"]
        y_offset = 0

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            match = regex.match(r"\$bandtype\{(\w+)\}", text)
            if match:
                self.current_band_type = match.group(1).lower()
                current_section = self.bands.get(self.current_band_type, self.bands["detail"])
                y_offset = 0
                continue

            if "$table{" in text:
                continue

            band = ET.SubElement(current_section, "band")
            band.set("height", "20")
            y_offset = self.add_paragraph(para, band, y_offset)

        detail_band = self.bands["detail"]

        for idx, table in enumerate(self.doc.tables[:1]):
            table_band = ET.SubElement(detail_band, "band")
            table_band.set("height", "100")
            table_height = self._convert_table_simple(table, table_band, 0)
            if table_height:
                table_band.set("height", str(table_height))

    def save(self, output_path):
        tree = ET.ElementTree(self.jrxml)
        with open(output_path, 'wb') as f:
            tree.write(f, pretty_print=True, xml_declaration=True, encoding="UTF-8")
        print(f"Diseño exportado a: {output_path}")
        print(f"Imágenes extraídas a: {self.images_dir}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python word_to_jasper.py 'archivo.docx'")
    else:
        input_file = sys.argv[1]
        output_file = input_file.replace(".docx", ".jrxml")
        conv = JasperConverter(input_file)
        conv.process()
        conv.save(output_file)